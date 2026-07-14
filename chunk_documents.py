import os
import json
from docx import Document

DOCS_FOLDER = "docs"
OUTPUT_FILE = "chunks.json"
CHUNK_SIZE = 400        # approx words per chunk
CHUNK_OVERLAP = 60      # words repeated between chunks, so context isn't lost at the edges

def read_docx(filepath):
    """Extract all non-empty paragraph text from a .docx file, including table cells."""
    doc = Document(filepath)
    full_text = []

    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                full_text.append(row_text)

    return "\n".join(full_text)

def chunk_text(text, source_name, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    """Split text into overlapping word-based chunks."""
    words = text.split()
    chunks = []
    start = 0

    while start < len(words):
        end = start + chunk_size
        chunk_words = words[start:end]
        chunk_text_str = " ".join(chunk_words)
        chunks.append({
            "source": source_name,
            "text": chunk_text_str
        })
        start += (chunk_size - overlap)

    return chunks

def main():
    all_chunks = []
    docx_files = [f for f in os.listdir(DOCS_FOLDER) if f.endswith(".docx")]

    if not docx_files:
        print(f"No .docx files found in '{DOCS_FOLDER}' folder.")
        return

    for filename in docx_files:
        filepath = os.path.join(DOCS_FOLDER, filename)
        print(f"Reading: {filename}")
        text = read_docx(filepath)
        chunks = chunk_text(text, source_name=filename)
        print(f"  -> {len(chunks)} chunks created")
        all_chunks.extend(chunks)

    with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
        json.dump(all_chunks, f, indent=2)

    print(f"\nDone. {len(all_chunks)} total chunks saved to {OUTPUT_FILE}")

if __name__ == "__main__":
    main()